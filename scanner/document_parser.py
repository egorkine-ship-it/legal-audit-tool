"""
Разбор документов сайта: извлечение текста из HTML/PDF/DOCX и вспомогательные
детекторы (формат, дата, версия).

Тяжёлые/опциональные зависимости (pypdf, python-docx, bs4) импортируются
внутри функций и обёрнуты в try/except, чтобы модуль импортировался даже без
этих библиотек. Публичные функции никогда не бросают исключения — при ошибке
возвращается безопасное значение.
"""
from __future__ import annotations

import io
import re
from typing import Tuple

from scanner import utils


# ---------------------------------------------------------------------------
# Извлечение текста из бинарных форматов
# ---------------------------------------------------------------------------
def parse_pdf(content: bytes) -> Tuple[str, str]:
    """Извлечь текст из PDF (bytes). Возвращает (text, error).

    Использует pypdf (опционально). Если библиотека недоступна или текст не
    извлекается — возвращает ("", "<код ошибки>"). Никогда не бросает.
    """
    if not content:
        return "", "empty_content"
    try:
        try:
            import pypdf  # type: ignore
        except ImportError:
            try:
                import PyPDF2 as pypdf  # type: ignore
            except ImportError:
                return "", "pypdf_not_installed"

        reader = None
        # pypdf >= 3 / PyPDF2 >= 2 используют PdfReader.
        try:
            reader = pypdf.PdfReader(io.BytesIO(content))
        except Exception:
            reader = None
        if reader is None:
            return "", "text_extraction_failed"

        parts = []
        try:
            pages = reader.pages
        except Exception:
            pages = []
        for page in pages:
            try:
                txt = page.extract_text() or ""
            except Exception:
                txt = ""
            if txt:
                parts.append(txt)
        text = "\n".join(parts).strip()
        if not text:
            return "", "text_extraction_failed"
        return text, ""
    except Exception as exc:  # предохранитель
        return "", f"pdf_error: {exc}"


def parse_docx(content: bytes) -> Tuple[str, str]:
    """Извлечь текст из DOCX (bytes) через python-docx. Возвращает (text, error)."""
    if not content:
        return "", "empty_content"
    try:
        try:
            import docx  # type: ignore
        except ImportError:
            return "", "python_docx_not_installed"

        try:
            document = docx.Document(io.BytesIO(content))
        except Exception:
            return "", "text_extraction_failed"

        parts = []
        try:
            for para in document.paragraphs:
                if para.text:
                    parts.append(para.text)
        except Exception:
            pass
        # Текст из таблиц.
        try:
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            parts.append(cell.text)
        except Exception:
            pass

        text = "\n".join(parts).strip()
        if not text:
            return "", "text_extraction_failed"
        return text, ""
    except Exception as exc:
        return "", f"docx_error: {exc}"


# ---------------------------------------------------------------------------
# HTML -> текст
# ---------------------------------------------------------------------------
_TAG_RE = re.compile(r"<[^>]+>")
_SCRIPT_STYLE_RE = re.compile(
    r"<(script|style|noscript|template)\b[^>]*>.*?</\1>",
    re.IGNORECASE | re.DOTALL,
)
_COMMENT_RE = re.compile(r"<!--.*?-->", re.DOTALL)


def html_to_text(html: str) -> str:
    """Преобразовать HTML в очищенный видимый текст.

    Использует bs4, если доступен; иначе — регулярный фолбэк со снятием тегов.
    Никогда не бросает.
    """
    if not html:
        return ""
    try:
        try:
            from bs4 import BeautifulSoup  # type: ignore
        except Exception:
            BeautifulSoup = None  # type: ignore

        if BeautifulSoup is not None:
            try:
                soup = BeautifulSoup(html, "html.parser")
                for tag in soup(["script", "style", "noscript", "template"]):
                    try:
                        tag.decompose()
                    except Exception:
                        pass
                text = soup.get_text(" ")
                return utils.clean_text(text)
            except Exception:
                pass  # падаем в регулярный фолбэк

        # Регулярный фолбэк.
        stripped = _COMMENT_RE.sub(" ", html)
        stripped = _SCRIPT_STYLE_RE.sub(" ", stripped)
        stripped = _TAG_RE.sub(" ", stripped)
        stripped = _unescape_basic(stripped)
        return utils.clean_text(stripped)
    except Exception:
        return ""


def _unescape_basic(text: str) -> str:
    try:
        import html as _html

        return _html.unescape(text)
    except Exception:
        # Минимальная замена самых частых сущностей.
        for a, b in (
            ("&nbsp;", " "),
            ("&amp;", "&"),
            ("&lt;", "<"),
            ("&gt;", ">"),
            ("&quot;", '"'),
            ("&#39;", "'"),
        ):
            text = text.replace(a, b)
        return text


# ---------------------------------------------------------------------------
# Диспетчер извлечения текста
# ---------------------------------------------------------------------------
def extract_text_from_bytes(content: bytes, fmt: str) -> Tuple[str, str]:
    """Извлечь текст из bytes согласно формату fmt in {html, pdf, docx}.

    Возвращает (text, error). Для html content декодируется best-effort.
    Никогда не бросает.
    """
    fmt = (fmt or "").lower().strip()
    try:
        if fmt == "pdf":
            return parse_pdf(content)
        if fmt == "docx":
            return parse_docx(content)
        if fmt == "html":
            if not content:
                return "", "empty_content"
            html = _decode_html_bytes(content)
            text = html_to_text(html)
            if not text:
                return "", "text_extraction_failed"
            return text, ""
        # Неизвестный формат — пробуем как обычный текст.
        if not content:
            return "", "empty_content"
        text = _decode_html_bytes(content).strip()
        if not text:
            return "", "text_extraction_failed"
        return text, ""
    except Exception as exc:
        return "", f"extract_error: {exc}"


def _decode_html_bytes(content: bytes) -> str:
    if not content:
        return ""
    # Попытка определить charset из meta.
    charset = ""
    try:
        head = content[:4096].decode("latin-1", errors="ignore")
        m = re.search(r"charset=[\"']?([\w\-]+)", head, re.IGNORECASE)
        if m:
            charset = m.group(1)
    except Exception:
        charset = ""
    for enc in (charset, "utf-8", "windows-1251", "cp1251", "latin-1"):
        if not enc:
            continue
        try:
            return content.decode(enc, errors="strict")
        except Exception:
            continue
    try:
        return content.decode("utf-8", errors="replace")
    except Exception:
        return ""


# ---------------------------------------------------------------------------
# Определение формата документа
# ---------------------------------------------------------------------------
def detect_format(url: str, content_type: str) -> str:
    """Определить формат документа по URL и Content-Type. Возвращает html/pdf/docx."""
    ct = (content_type or "").lower()
    ext = ""
    try:
        ext = utils.url_extension(url or "")
    except Exception:
        ext = ""

    # По расширению — самый надёжный признак.
    if ext == "pdf" or "pdf" in ct:
        return "pdf"
    if ext in ("docx", "doc") or "officedocument.wordprocessing" in ct or "msword" in ct:
        return "docx"
    if ext in ("htm", "html") or "html" in ct or "xml" in ct or "text/plain" in ct:
        return "html"

    # По умолчанию считаем документ HTML-страницей.
    return "html"


# ---------------------------------------------------------------------------
# Детекторы даты и версии в тексте документа
# ---------------------------------------------------------------------------
def detect_date(text: str) -> str:
    """Первая обнаруженная дата в тексте (или "")."""
    try:
        dates = utils.find_dates(text or "")
        return dates[0] if dates else ""
    except Exception:
        return ""


def detect_version(text: str) -> str:
    """Номер версии/редакции документа (или "")."""
    if not text:
        return ""
    try:
        m = utils.RE_VERSION.search(text)
        if m:
            return (m.group(1) or "").strip()
        return ""
    except Exception:
        return ""
